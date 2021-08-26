import base64

import djangoProject.settings
import json
from io import BytesIO

import pytz
# Create your views here.
from .forms import *
import datetime
from Qn.form import *
from Qn.models import *
from .export import *
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

utc = pytz.UTC

IS_LINUX = False
try:
    import pythoncom
except:
    IS_LINUX = True


# Create your views here.

@csrf_exempt
def empty_the_recycle_bin(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        username_form = UserNameForm(request.POST)
        if username_form.is_valid():
            username = username_form.cleaned_data.get('username')
            this_username = request.session.get('username')
            if this_username != username:
                return JsonResponse({'status_code': 0})
            qn_list = Survey.objects.filter(username=username, is_deleted=True)
            for qn in qn_list:
                qn.delete()
            return JsonResponse(response)
        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def delete_survey_not_real(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                survey = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': -1, 'message': '问卷不存在'}
                return JsonResponse(response)
            if survey.is_deleted == True:
                response = {'status_code': 0, 'message': '问卷已放入回收站'}
                return JsonResponse(response)
            survey.is_deleted = True

            survey.is_released = False
            survey.is_finished = False
            survey.is_collected = False
            survey.save()
            return JsonResponse(response)
        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def recover_survey_from_delete(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                survey = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': -1, 'message': '问卷不存在'}
                return JsonResponse(response)
            if survey.is_deleted == False:
                response = {'status_code': 3, 'message': '问卷未放入回收站'}
                return JsonResponse(response)

            survey.is_deleted = False
            survey.is_released = False
            survey.is_finished = False
            survey.is_collected = False
            survey.save()
            return JsonResponse(response)
        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


def produce_time(example):
    if example is None or example == '':
        return False
    else:
        return True


def get_qn_data(qn_id):
    id = qn_id
    survey = Survey.objects.get(survey_id=qn_id)
    response = {'status_code': 1, 'message': 'success'}
    response['qn_id'] = survey.survey_id
    response['username'] = survey.username
    response['title'] = survey.title
    response['description'] = survey.description
    response['type'] = survey.type

    response['question_num'] = survey.question_num
    response['created_time'] = response['release_time'] = response['finished_time'] = ''
    if produce_time(survey.created_time):
        response['created_time'] = survey.created_time.strftime("%Y/%m/%d %H:%M")
    if produce_time(survey.finished_time):
        response['finished_time'] = survey.finished_time.strftime("%Y/%m/%d %H:%M")
    if produce_time(survey.release_time):
        response['release_time'] = survey.release_time.strftime("%Y/%m/%d %H:%M")
    response['is_released'] = survey.is_released
    response['is_deleted'] = survey.is_deleted
    response['is_finished'] = survey.is_finished
    response['share_url'] = survey.share_url
    response['docx_url'] = survey.docx_url
    response['pdf_url'] = survey.pdf_url
    response['excel_url'] = survey.excel_url
    response['recycling_num'] = survey.recycling_num
    response['max_recycling'] = survey.max_recycling


    question_list = Question.objects.filter(survey_id=qn_id)
    questions = []
    for item in question_list:
        temp = {}
        temp['question_id'] = item.question_id
        temp['row'] = item.raw
        temp['score'] = item.score
        temp['title'] = item.title
        temp['description'] = item.direction
        temp['must'] = item.is_must_answer
        temp['type'] = item.type
        temp['qn_id'] = qn_id
        temp['sequence'] = item.sequence
        temp['option_num'] = item.option_num
        temp['refer'] = item.right_answer
        temp['point'] = item.point
        temp['id'] = item.sequence  # 按照前端的题目顺序
        temp['options'] = []
        temp['answer'] = item.right_answer
        if temp['type'] in ['radio', 'checkbox']:
            # 单选题或者多选题有选项
            option_list = Option.objects.filter(question_id=item.question_id)
            for option_item in option_list:
                option_dict = {}
                option_dict['id'] = option_item.option_id
                option_dict['title'] = option_item.content
                temp['options'].append(option_dict)

        else:  # TODO 填空题或者其他
            pass

        questions.append(temp)
        print(questions)
    response['questions'] = questions
    return response


@csrf_exempt
def delete_survey_real(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                survey = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': -1, 'message': '问卷不存在'}
                return JsonResponse(response)

            username = request.session.get('username')
            if survey.username != username:
                return JsonResponse({'status_code': 0})

            survey.delete()
            # 是否真的删掉呢
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def get_survey_details(request):
    response = {'status_code': 1, 'message': 'success'}
    this_username = request.session.get('username')
    # if not this_username:
    #     return JsonResponse({'status_code': 0})
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                survey = Survey.objects.get(survey_id=id)
                print(survey.survey_id)
            except:
                response = {'status_code': -2, 'message': '问卷不存在'}
                return JsonResponse(response)

            # if survey.username != this_username:
            #     return JsonResponse({'status_code': 0})

            response = get_qn_data(id)

            return JsonResponse(response)
        else:
            response = {'status_code': -1, 'message': '问卷id不为整数'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def get_survey_details_by_others(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        code = request.POST.get('code')
        try:
            print(code)
            survey = Survey.objects.get(share_url=code)
        except:
            response = {'status_code': 2, 'message': '问卷不存在'}
            return JsonResponse(response)
        if survey.is_released and not survey.is_deleted:
            response = get_qn_data(survey.survey_id)
            return JsonResponse(response)
        else:
            return JsonResponse({'status_code': 3})
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def delete_question(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        question_form = QuestionIdForm(request.POST)
        if question_form.is_valid():
            id = question_form.cleaned_data.get('question_id')
            try:
                question = Question.objects.get(question_id=id)
            except:
                response = {'status_code': -1, 'message': '题目不存在'}
                return JsonResponse(response)
            question.delete()
            # 是否真的删掉呢
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def delete_option(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        option_form = OptionIdForm(request.POST)
        if option_form.is_valid():
            id = option_form.cleaned_data.get('option_id')
            try:
                option = Option.objects.get(option_id=id)
            except:
                response = {'status_code': -1, 'message': '选项不存在'}
                return JsonResponse(response)
            option.delete()
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


# username title description type
@csrf_exempt
def create_qn(request):
    global survey
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        new_qn_form = CreateNewQnForm(request.POST)
        if new_qn_form.is_valid():
            username = new_qn_form.cleaned_data.get('username')
            title = new_qn_form.cleaned_data.get('title')
            description = new_qn_form.cleaned_data.get('description')
            type = new_qn_form.cleaned_data.get('type')

            description = "这里是问卷说明信息，您可以在此处编写关于本问卷的简介，帮助填写者了解这份问卷。"
            if type == '2':
                description = "这里是考试问卷说明信息，您可以在此处编写关于本考试问卷的简介，帮助填写者了解这份问卷。"

            try:
                user = User.objects.get(username=username)

            except:
                response = {'status_code': 2, 'message': '用户不存在'}
                return JsonResponse(response)

            # if request.session.get('username') != username:
            #     return JsonResponse({'status_code': 2})

            if title == '':
                title = "默认标题"
                if type == '2':
                    title = "默认标题"

            try:
                survey = Survey(username=username, title=title, type=type, description=description, question_num=0,
                                recycling_num=0)
                survey.save()
            except:
                response = {'status_code': -3, 'message': '后端炸了'}
                return JsonResponse(response)

            response['qn_id'] = survey.survey_id
            return JsonResponse(response)


        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': 'invalid http method'}
        return JsonResponse(response)


@csrf_exempt
def create_option(question, content, sequence):
    option = Option()
    option.content = content
    question.option_num += 1
    option.question_id = question
    question.save()
    option.order = sequence
    option.save()


#  title direction is_must_answer type qn_id options:只传option的title字符串使用特殊字符例如 ^%之类的隔开便于传输
@csrf_exempt
def create_question(request):
    # TODO 完善json。dump
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        new_question_form = CreateNewQuestionForm(request.POST)
        if new_question_form.is_valid():
            question = Question()
            try:
                question.title = new_question_form.cleaned_data.get('title')
                question.direction = new_question_form.cleaned_data.get('description')
                question.is_must_answer = new_question_form.cleaned_data.get('must')
                question.type = new_question_form.cleaned_data.get('type')
                survey_id = new_question_form.cleaned_data.get('qn_id')
                question.survey_id = Survey.objects.get(survey_id=survey_id)
                question.raw = new_question_form.cleaned_data.get('row')
                question.score = new_question_form.cleaned_data.get('score')

                option_str = new_question_form.cleaned_data.get('options')
            except:
                response = {'status_code': -3, 'message': '后端炸了'}
                return JsonResponse(response)
            KEY = "^_^_^"
            option_list = option_str.split(KEY)
            for item in option_list:
                create_option(question, item)
            question.save()
            response['option_num'] = len(option_list)

            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)

    else:
        response = {'status_code': -2, 'message': 'invalid http method'}
        return JsonResponse(response)


@csrf_exempt
def save_qn(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        req = json.loads(request.body)
        print(req)
        qn_id = req['qn_id']
        try:
            question_list = Question.objects.filter(survey_id=qn_id)
        except:
            response = {'status_code': 3, 'message': '问卷不存在'}
            return JsonResponse(response)
        for question in question_list:
            question.delete()

        survey = Survey.objects.get(survey_id=qn_id)
        survey.username = req['username']
        survey.title = req['title']
        survey.description = req['description']
        survey.type = req['type']
        survey.save()
        question_list = req['questions']

        if request.session.get("username") != req['username']:
            request.session.flush()
            return JsonResponse({'status_code': 0})

        # TODO
        question_num = 0
        for question in question_list:
            question_num += 1
            create_question_in_save(question['title'], question['description'], question['must']
                                    , question['type'], qn_id=req['qn_id'], raw=question['row'],
                                    score=question['score'],
                                    options=question['options'],
                                    sequence=question['id'],
                                    )

        survey.question_num = question_num
        survey.save()
        return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': 'invalid http method'}


def create_question_in_save(title, direction, must, type, qn_id, raw, score, options, sequence,refer ,point):
    question = Question()
    try:
        question.title = title
        question.direction = direction
        question.is_must_answer = must
        question.type = type
        survey_id = qn_id
        question.survey_id = Survey.objects.get(survey_id=survey_id)
        question.raw = raw
        question.score = score
        question.sequence = sequence
        question.point = point
        question.right_answer =refer
    except:
        response = {'status_code': -3, 'message': '后端炸了'}
        return JsonResponse(response)
    KEY = "^_^_^"

    option_list = options
    for item in option_list:
        print(item)
        content = item['title']
        sequence = item['id']
        create_option(question, content, sequence)
    question.save()


@csrf_exempt
def deploy_qn(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                survey = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)

            if survey.is_deleted == True:
                response = {'status_code': 4, 'message': '问卷已经放入回收站'}
                return JsonResponse(response)

            if survey.is_released == True:
                response = {'status_code': 3, 'message': '问卷已经发布，不要重复操作'}
                return JsonResponse(response)
            if survey.is_finished == True:
                response = {'status_code': 6, 'message': '问卷已经结束，不可操作'}
                return JsonResponse(response)
            if (Question.objects.filter(survey_id=survey)).count() == 0:
                response = {'status_code': 7, 'message': '问卷中无问题，不可发行'}
                return JsonResponse(response)

            survey.is_released = True
            if survey.release_time == '' or survey.release_time is None or survey.release_time == 'null':
                survey.release_time = datetime.datetime.now()
            if survey.share_url == '' or survey.share_url is None or survey.share_url == 'null':
                # 生成问卷码
                code = hash_code(survey.username, str(survey.survey_id))
                # code = hash_code(code, datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                end_info = code[:20].upper()
                while Survey.objects.filter(share_url=end_info):
                    code = hash_code(code, datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    end_info = code[:20].upper()

                survey.share_url = end_info
                survey.save()
                return JsonResponse({'status_code': 10, 'code': survey.share_url})
            else:
                survey.save()
                return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def pause_qn(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                survey = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': -1, 'message': '问卷不存在'}
                return JsonResponse(response)
            if survey.is_deleted == True:
                response = {'status_code': 2, 'message': '问卷已放入回收站'}
                return JsonResponse(response)
            if survey.is_finished == True:
                response = {'status_code': 6, 'message': '问卷已经结束，不可操作'}
                return JsonResponse(response)
            if not survey.is_released:
                response = {'status_code': 3, 'message': '问卷为发行，不可取消发行'}
                return JsonResponse(response)
            survey.is_released = False
            survey.save()
            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def finish_qn(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                survey = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)
            if survey.is_deleted == True:
                response = {'status_code': 4, 'message': '问卷已经放入回收站'}
                return JsonResponse(response)
            if survey.is_finished:
                response = {'status_code': 5, 'message': '问卷已经停止回收'}
                return JsonResponse(response)
            survey.is_finished = True
            survey.finished_time = datetime.datetime.now()
            survey.is_released = False

            survey.save()
            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx import *


@csrf_exempt
def TestDocument(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                survey = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)
            if survey.type == '2':
                document, f, docx_title, _ = paper_to_docx(id)

            else: # TODO
                document, f, docx_title, _ = qn_to_docx(id)

            response['filename'] = docx_title
            response['docx_url'] = djangoProject.settings.WEB_ROOT + "/media/Document/" + docx_title
            # TODO: 根据实时文件位置设置url
            survey.docx_url = response['docx_url']
            survey.save()
            response['b64data'] = base64.b64encode(f.getvalue()).decode()
            # print(f.getvalue())
            # print(response['Content-Length'])

            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


# 根据问卷id传递文件格式返回上一个函数。具体正在写。
# 只要文件能打开就好写了
import hashlib


def hash_code(s, salt='Qn'):  # generate s+salt into hash_code (default: salt=online publish)
    h = hashlib.sha256()
    s += salt
    h.update(s.encode())  # update method get bytes(type)
    return h.hexdigest()


def qn_to_docx(qn_id):
    document = Document()
    survey = Survey.objects.get(survey_id=qn_id)
    docx_title = survey.title + '_' + str(survey.username) + '_' + str(qn_id) + ".docx"

    # code = hash_code(str(survey.username),str(qn_id))

    # docx_title = code
    print(docx_title)

    # run = document.add_paragraph().add_run('This is a letter.')
    # font = run.font
    # font.name = '宋体' 英文字体设置
    document.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER).font.name = '宋体'  # 添加字体样式-Song
    document.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # document.add_paragraph().add_run('第二个段落，abcDEFg，这是中文字符', style='Song')

    document.add_heading(survey.title, 0)

    paragraph_list = []

    paragraph = document.add_paragraph().add_run(survey.description, style='Song')

    introduction = "本问卷已经收集了" + str(survey.recycling_num) + "份，共计" + str(survey.question_num) + "个问题"
    paragraph = document.add_paragraph().add_run(introduction, style='Song')
    paragraph_list.append(paragraph)

    questions = Question.objects.filter(survey_id=survey)
    i = 1
    for question in questions:

        type = question.type
        type_str = ""
        if type == 'radio':
            type_str = "单选题"
        elif type == 'checkbox':
            type_str = '多选题'
        elif type == 'text':
            type_str = '填空题'
        elif type == 'mark':
            type_str = '评分题'
        document.add_paragraph().add_run(str(i) + "、" + question.title + "(" + type_str + ")", style='Song')

        i += 1
        options = Option.objects.filter(question_id=question)
        option_option = 0
        num = 1
        for option in options:
            option_str = "      "

            alphas = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

            if question.type in ['checkbox', 'radio']:
                # option_str += alphas[option_option] + " :  "
                option_str += "选项 " + str(num) + " :  "
                option_option += 1
                num += 1

            option_str += option.content
            document.add_paragraph().add_run(option_str, style='Song')
        if question.type in ['mark', 'text']:
            document.add_paragraph(' ')

    document.add_page_break()
    # document.add_paragraph(str(qn_id))
    f = BytesIO()
    save_path = docx_title

    document.save(f)
    # document.save(save_path)

    docx_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    if IS_LINUX:
        docx_path = djangoProject.settings.MEDIA_ROOT + "/Document/"

    print(docx_path)
    document.save(docx_path + docx_title)

    return document, f, docx_title, docx_path


from docx2pdf import convert


def qn_to_pdf(qn_id):
    qn = Survey.objects.get(survey_id=qn_id)
    if qn.type == '2':
        document, _, docx_title, docx_path = paper_to_docx(id)
    else:
        document, _, docx_title, docx_path = qn_to_docx(qn_id)
    input_file = docx_path + docx_title
    out_file = docx_path + docx_title.replace('.docx', '.pdf')
    pdf_title = docx_title.replace('.docx', '.pdf')
    try:
        import pythoncom
        pythoncom.CoInitialize()
        convert(input_file, out_file)
    except:
        doc2pdf_linux(input_file, docx_path)

    return pdf_title


@csrf_exempt
def pdf_document(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                qn = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)

            pdf_title = qn_to_pdf(qn.survey_id)
            response['filename'] = pdf_title
            response['pdf_url'] = djangoProject.settings.WEB_ROOT + "/media/Document/" + pdf_title
            # TODO: 根据实时文件位置设置url
            qn.pdf_url = response['pdf_url']
            qn.save()

            return JsonResponse(response)


        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


import xlwt

from Qn.views import KEY_STR
def write_submit_to_excel(qn_id):
    qn = Survey.objects.get(survey_id=qn_id)
    submit_list = Submit.objects.filter(survey_id=qn)

    xls = xlwt.Workbook()
    sht1 = xls.add_sheet("Sheet1")

    sht1.write(0, 0, "序号")
    sht1.write(0, 1, "提交者")
    sht1.write(0, 2, "提交时间")
    question_list = Question.objects.filter(survey_id=qn)
    question_num = len(question_list)
    i = 1

    for question in question_list:
        sht1.write(0, 2 + i, str(i) + "、" + question.title)
        i += 1

    id = 1
    for submit in submit_list:
        sht1.write(id, 0, id)
        username = submit.username
        if username == '' or username is None:
            username = "匿名用户"
        sht1.write(id, 1, username)
        sht1.write(id, 2, submit.submit_time.strftime("%Y/%m/%d %H:%M"))
        question_num = 1
        for question in question_list:
            answer_str = ""
            try:
                answer = Answer.objects.get(submit_id=submit, question_id=question)
                answer_str = answer.answer
            except:
                answer_str = ""
            if question.type == 'checkbox':
                answer_str = answer_str.replace(KEY_STR,';')

            sht1.write(id, 2 + question_num, answer_str)

            question_num += 1

        id += 1
    save_path = djangoProject.settings.MEDIA_ROOT + "\Document\\"
    if IS_LINUX:
        save_path = djangoProject.settings.MEDIA_ROOT + "/Document/"
    excel_name = qn.title + "问卷的统计信息" + ".xls"
    xls.save(save_path + excel_name)
    return excel_name


@csrf_exempt
def export_excel(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                qn = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)
            username = qn.username
            # if request.session['username'] != username:
            #     response = {'status_code': 0, 'message': '没有访问权限'}
            #     return JsonResponse(response)
            try:
                submit_list = Submit.objects.filter(survey_id=qn)
                # 找不到问卷提交
            except():
                response = {'status_code': 3, 'message': '该问卷暂无提交，无法导出'}
                return JsonResponse(response)
            if len(submit_list) == 0:
                response = {'status_code': 3, 'message': '该问卷暂无提交，无法导出'}
                return JsonResponse(response)

            excel_name = write_submit_to_excel(id)

            response['excel_url'] = djangoProject.settings.WEB_ROOT + "/media/Document/" + excel_name
            qn.excel_url = response['excel_url']
            response['excel_name'] = excel_name

            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def duplicate_qn(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                qn = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)
            new_qn = Survey(title=qn.title+"-副本", description=qn.description, question_num=qn.question_num, recycling_num=0,
                            username=qn.username, type=qn.type)

            new_qn.save()
            new_qn_id = new_qn.survey_id
            questions = Question.objects.filter(survey_id=qn)
            for question in questions:
                new_question = Question(title=question.title, direction=question.direction,
                                        is_must_answer=question.is_must_answer,
                                        sequence=question.sequence, option_num=question.option_num,
                                        score=question.score, raw=question.raw,
                                        type=question.type, survey_id=new_qn)
                new_question.save()
                options = Option.objects.filter(question_id=question)

                for option in options:
                    new_option = Option(content=option.content, question_id=new_question, order=option.order)
                    new_option.save()

            print(new_qn_id)
            # response = get_qn_data(new_qn_id)
            return JsonResponse({'status_code': 1, 'qn_id': new_qn_id})

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def empty_qn_all_Submit(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                qn = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)
            username = qn.username
            if request.session['username'] != username:
                response = {'status_code': 0, 'message': '没有访问权限'}
                return JsonResponse(response)

            submit_list = Submit.objects.filter(survey_id=qn.survey_id)
            for submit in submit_list:
                submit.delete()
            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


import subprocess


def doc2pdf_linux(docPath, pdfPath):
    """
    convert a doc/docx document to pdf format (linux only, requires libreoffice)
    :param doc: path to document
    """
    cmd = 'libreoffice7.0 --headless --invisible  --convert-to pdf:writer_pdf_Export'.split() + [docPath] + [
        '--outdir'] + [pdfPath]
    print(cmd)
    p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    p.wait(timeout=30)
    stdout, stderr = p.communicate()
    if stderr:
        raise subprocess.SubprocessError(stderr)


# def doc2pdf_linux(docPath, pdfPath):
#     """
#     convert a doc/docx document to pdf format (linux only, requires libreoffice)
#     :param doc: path to document
#     """
#     cmd = 'libreoffice7.0  --headless --convert-to pdf'.split() + [docPath] + ['--outdir'] + [pdfPath]
#     p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
#     p.wait(timeout=30)
#     stdout, stderr = p.communicate()
#     if stderr:
#         raise subprocess.SubprocessError(stderr)

def question_dict_to_question(question, question_dict):
    # question = Question()#TODO delete
    question.title = question_dict['title']
    print(question_dict)
    question.direction = question_dict['description']
    question.is_must_answer = question_dict['must']
    question.type = question_dict['type']
    # qn_id =  question_dict['qn_id']
    # question.survey_id = Survey.objects.get(survey_id=qn_id)
    question.raw = question_dict['row']
    question.score = question_dict['score']
    options = question_dict['options']
    question.sequence = question_dict['id']

    option_list_delete = Option.objects.filter(question_id=question)
    for option in option_list_delete:
        option.delete()
    option_list = options
    # option_list = option_set[0]
    # print("options") ,print(options)
    print(option_list)

    for item in option_list:
        print("item", end=" ")
        print(item)
        content = item['title']
        sequence = item['id']
        create_option(question, content, sequence)
    question.save()


@csrf_exempt
def save_qn_keep_history(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        req = json.loads(request.body)
        print(req)
        qn_id = req['qn_id']
        try:
            questions = Question.objects.filter(survey_id=qn_id)
        except:
            response = {'status_code': 3, 'message': '问卷不存在'}
            return JsonResponse(response)
        submit_list =  Survey.objects.filter(survey_id=qn_id)
        for submit in submit_list:
            submit.is_valid = False
            submit.save()

        survey = Survey.objects.get(survey_id=qn_id)
        survey.username = req['username']
        survey.title = req['title']
        if survey.title == '':
            survey.title = "默认标题"
        survey.description = req['description']
        if survey.description == '':
            survey.description = "这里是问卷说明信息，您可以在此处编写关于本问卷的简介，帮助填写者了解这份问卷。"
        survey.type = req['type']
        if req['type'] == '2':
            # 如果问卷是考试问卷
            #TODO 正常发问卷的截止时间
            survey.finished_time = req['finished_time']

        survey.save()
        question_list = req['questions']



        # if request.session.get("username") != req['username']:
        #     request.session.flush()
        #     return JsonResponse({'status_code': 0})

        for question in questions:
            num = 0
            for question_dict in question_list:
                if question_dict['question_id'] == question.question_id:
                    #旧问题在新问题中有 更新问题
                    question_dict_to_question(question,question_dict)
                    num = 1
                    break

            if num == 0:
                question.delete()
        for question_dict in question_list:
            # num = 1
            # for question in questions:
            #     if question_dict['question_id'] != 0:
            #         break
            #     num += 1
            # if num == len(questions):
            try:
                question_dict['question_id'] = question_dict['question_id']
            except:
                question_dict['question_id'] = 0
            refer = ''
            point = 0
            if req['type'] == '2':
                refer = question_dict['refer']
                point = question_dict['point']

            if question_dict['question_id'] == 0:
                create_question_in_save(question_dict['title'], question_dict['description'], question_dict['must']
                                        , question_dict['type'], qn_id=req['qn_id'], raw=question_dict['row'],
                                        score=question_dict['score'],
                                        options=question_dict['options'],
                                        sequence=question_dict['id'],refer=refer,point=point
                                        )
                # 添加问题


        question_num = 0

        survey.save()
        question_list = Question.objects.filter(survey_id=survey)
        for question in question_list:
            question_num += 1
        survey.question_num = question_num
        print("保存成功，该问卷的问题数目为："+str(question_num))
        survey.save()

        return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': 'invalid http method'}
        return JsonResponse(response)

# def get_question_detail(question):
#     question =Question()
#     #TODO delete question
#     dic = {}
#     dic['id'] = question.sequence
#     dic['sequence'] = question.sequence
#     dic['title'] = question.title
#     dic['description'] = question.direction


@csrf_exempt
def get_answer_from_submit(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        # survey_form = SurveyIdForm(request.POST)
        submit_form = SubmitIDForm(request.POST)
        # print(submit_form)
        if submit_form.is_valid():
            id = submit_form.cleaned_data.get('submit_id')
            print(id)
            try:
                submit = Submit.objects.get(submit_id=id)
            except:
                response = {'status_code': 2, 'message': '答卷不存在'}
                return JsonResponse(response)
            qn = submit.survey_id
            qn_dict = get_qn_data(qn.survey_id)
            questions = qn_dict['questions']
            response['questions'] = questions
            # TODO
            # if request.session['username'] != username:
            #     response = {'status_code': 0, 'message': '没有访问权限'}
            #     return JsonResponse(response)

            try:
                answer_list = Answer.objects.filter(submit_id=submit)
                if len(answer_list) == 0:
                    raise Exception('')
            except:
                response = {'status_code': 3, 'message': '该问卷暂无回答'}
                return JsonResponse(response)
            answers = []
            response['submit_id'] = submit.submit_id
            response['submit_time'] = submit.submit_time.strftime("%Y/%m/%d %H:%M")
            response['username'] = submit.username
            response['is_valid'] = submit.is_valid
            response['score'] = submit.score
            response['qn_id'] = submit.survey_id.survey_id

            for answer in answer_list:
                item = {}
                item['answer'] = answer.answer
                item['score'] = answer.score
                item['username'] = answer.username
                item['answer_id'] = answer.answer_id
                item['type'] = answer.type
                item['question_id'] = answer.question_id.question_id
                item['submit_id'] = id
                answers.append(item)
            # TODO
            response['answers'] = answers
            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


# 当天回收，当周，总税收，前五天的返回时间返回日期
@csrf_exempt
def get_qn_recycling_num(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                qn = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)
            username = qn.username
            # TODO
            # if request.session['username'] != username:
            #     response = {'status_code': 0, 'message': '没有访问权限'}
            #     return JsonResponse(response)

            num_all = len(Submit.objects.filter(survey_id=qn))

            today_exact = datetime.datetime.now()
            today = datetime.datetime(year=today_exact.year, month=today_exact.month, day=today_exact.day)
            yesterday = today - datetime.timedelta(days=1)
            a_week_ago = today - datetime.timedelta(days=7)
            day_list = []
            for i in range(7, 0, -1):
                the_day = today - datetime.timedelta(days=i)
                print(the_day)
                day_list.append(the_day)
            print(today)
            num_day = len(Submit.objects.filter(survey_id=qn, submit_time__gte=today))
            num_week = len(Submit.objects.filter(survey_id=qn, submit_time__gte=a_week_ago))

            response['num_week'] = num_week
            response['num_day'] = num_day
            response['num_all'] = num_all

            dates = []
            nums = []
            for i in range(4, -1, -1):
                before = today - datetime.timedelta(days=i)
                after = today - datetime.timedelta(days=i - 1)
                num = len(Submit.objects.filter(survey_id=qn, submit_time__gte=before, submit_time__lte=after))
                nums.append(num)
                date_str = before.strftime("%m.%d")
                if date_str[0] == '0':
                    date_str = date_str[1:]
                dates.append(date_str)

            response['nums'] = nums
            response['dates'] = dates

            return JsonResponse(response)
        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def delete_submit(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        # survey_form = SurveyIdForm(request.POST)
        submit_form = SubmitIDForm(request.POST)
        # print(submit_form)
        if submit_form.is_valid():
            id = submit_form.cleaned_data.get('submit_id')
            print(id)
            try:
                submit = Submit.objects.get(submit_id=id)
            except:
                response = {'status_code': 2, 'message': '答卷不存在'}
                return JsonResponse(response)

            # TODO
            # if request.session['username'] != username:
            #     response = {'status_code': 0, 'message': '没有访问权限'}
            #     return JsonResponse(response)
            answer_list = Answer.objects.filter(submit_id=submit)
            # for answer in answer_list:
            #     answer.delete()
            # 数据库是级联删除的，删除了submit 带有这个外键的自动珊瑚
            submit.delete()

            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def get_qn_all_submit(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                qn = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)
            # username = qn.username
            # if request.session['username'] != username:
            #     response = {'status_code': 0, 'message': '没有访问权限'}
            #     return JsonResponse(response)

            question_sum = qn.question_num
            print(question_sum)
            submits = []
            submit_list = Submit.objects.filter(survey_id=qn.survey_id)
            i = 1
            for submit in submit_list:
                item = {}
                item['num'] = i
                i += 1
                item['submit_id'] = submit.submit_id
                item['submit_time'] = submit.submit_time.strftime("%Y/%m/%d %H:%M")
                if submit.username and submit.username != '':
                    item['username'] = submit.username
                else:
                    item['username'] = '匿名用户'
                item['is_valid'] = submit.is_valid
                item['score'] = submit.score
                item['qn_id'] = submit.survey_id.survey_id

                answer_num = len(Answer.objects.filter(submit_id=submit))
                item['answer_num'] = answer_num
                item['answer_percent'] = str(int((float(answer_num) / question_sum) * 100)) + '%'
                submits.append(item)

            response['submits'] = submits
            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def cross_analysis(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        cross_form = CrossAnalysisForm(request.POST)
        if cross_form.is_valid():
            question_id_1 = cross_form.cleaned_data.get('question_id_1')
            question_id_2 = cross_form.cleaned_data.get('question_id_2')
            try:
                question_1 = Question.objects.get(question_id=question_id_1)
                question_2 = Question.objects.get(question_id=question_id_2)
            except:
                response = {'status_code': 2, 'message': '问题不存在'}
                return JsonResponse(response)
            qn = question_1.survey_id
            num_list = [[int(0) for x in range(0, question_1.option_num + +7)] for y in
                        range(0, question_2.option_num + 7)]
            submit_list = Submit.objects.filter(survey_id=qn)
            option_list1 = Option.objects.filter(question_id=question_1)
            option_list2 = Option.objects.filter(question_id=question_2)
            for option in option_list1:
                print(option.content, end=" ")
            print()
            for option in option_list2:
                print(option.content, end=" ")

            # for submit in submit_list:
            #
            #     answer_list = Answer.objects.filter(submit_id=submit)
            #     i = 1; j = 1;

            # i=1
            # for option in option_list1:
            #     # answer_list = Answer.objects.all()
            #     answers = []
            #     for answer in Answer.objects.all():
            #         if answer.question_id.survey_id == question_1.survey_id:
            #             answers.append(answer)
            #     for answer in answers:
            #         if answer.answer.find(option.content) >= 0:
            #             # submit = answer.submit_id
            #             answer_q2_list = Answer.objects.filter(submit_id=answer.submit_id,question_id=question_2)
            #             for answer_q2 in answer_q2_list:
            #                 j = 1
            #                 for oprion_q2 in option_list2:
            #                     if answer_q2.answer.find(oprion_q2.content) >= 0:
            #                         num_list[i][j] += 1
            #                     j += 1
            #     i+=1

            submit_list = Submit.objects.filter(survey_id=qn)
            for submit in submit_list:
                answer1 = Answer.objects.get(submit_id=submit, question_id=question_1)
                answer2 = Answer.objects.get(submit_id=submit, question_id=question_2)
                selection_x = []
                selection_y = []

                i = 1
                for option in option_list1:
                    if answer1.answer.find(option.content) >= 0:
                        selection_x.append(i)
                    i += 1

                j = 1
                for option in option_list2:
                    if answer2.answer.find(option.content) >= 0:
                        selection_y.append(j)
                    j += 1
                for i in selection_x:
                    for j in selection_y:
                        num_list[i][j] += 1

            tableData = []
            tableHead = []
            item = {}
            item['column_name'] = "column_0"
            item['column_comment'] = "子问题"
            tableHead.append(item)
            j = 1
            for option in option_list2:
                item = {}
                item['column_name'] = "column_{}".format(j)
                item['column_comment'] = option.content
                tableHead.append(item)
                j += 1

            # tableData.append(item)
            i = 1
            for option in option_list1:
                item = {}
                sum = 0
                for kk in range(len(option_list2) + 1):
                    sum += num_list[i][kk]
                item['column_0'] = option.content
                for j in range(1, len(option_list2) + 1):
                    recycling_num = qn.recycling_num
                    if recycling_num == 0:
                        recycling_num = 1
                    if sum == 0:
                        ret = "0(0%)"
                    else:
                        ret = str(num_list[i][j]) + "(" + str(int(num_list[i][j] * 100 / sum)) + "%)"

                    item['column_{}'.format(j)] = ret
                i += 1
                tableData.append(item)
            response['tableHead'] = tableHead
            response['tableData'] = tableData
            response['num_list'] = num_list
            return JsonResponse(response)

        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def get_qn_question(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            try:
                qn = Survey.objects.get(survey_id=id)
            except:
                response = {'status_code': 2, 'message': '问卷不存在'}
                return JsonResponse(response)
            # username = qn.username
            # if request.session['username'] != username:
            #     response = {'status_code': 0, 'message': '没有访问权限'}
            #     return JsonResponse(response)

            question_list = Question.objects.filter(survey_id=qn)
            i = 1
            questions = []
            for question in question_list:

                if question.type in ['radio', 'checkbox']:
                    item = {}
                    item['value1'] = item['value2'] = i
                    i += 1
                    item['label'] = question.title
                    item['question_id'] = question.question_id
                    questions.append(item)
            response['questions'] = questions

            return JsonResponse(response)


        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)


@csrf_exempt
def submit_reporter(request):
    response = {'status_code': 1, 'message': 'success'}
    if request.method == 'POST':
        survey_form = SurveyIdForm(request.POST)
        if survey_form.is_valid():
            id = survey_form.cleaned_data.get('qn_id')
            print("用户请求查看问卷 " + str(id) + " 的数据")
            question_list = Question.objects.filter(survey_id=id)
            questions = []
            for question in question_list:
                item = {}
                item['id'] = question.sequence
                item['title'] = question.title
                item['type'] = question.type
                item['row'] = question.raw
                item['score'] = question.score
                item['must'] = question.is_must_answer
                answer_list = Answer.objects.filter(question_id=question)
                option_list = Option.objects.filter(question_id=question)
                option_contnet_list = []
                for option in option_list:
                    option_contnet_list.append(option.content)
                options = []
                if item['type'] in ['checkbox', 'radio']:
                    option_list = Option.objects.filter(question_id=question)
                    for option in option_list:
                        dict = {}
                        dict['id'] = option.order
                        dict['title'] = option.content
                        dict['choosed'] = 0
                        options.append(dict)

                    for answer in answer_list:
                        i=0
                        for option_title in option_contnet_list:
                            if answer.answer.find(option_title)>=0:
                                options[i]['choosed'] += 1
                            i += 1


                elif item['type'] == 'text':
                    tableData = []
                    num = 1
                    for answer in answer_list:
                        dict = {}
                        dict['num'] = num
                        dict['answer'] = answer.answer
                        tableData.append(dict)
                        num += 1
                    item['tableData'] = tableData

                elif item['type'] == 'mark':
                    options= []

                    for i in range(1,question.score+1):
                        dict = {}
                        dict['title'] = i
                        dict['choosed'] = 0
                        options.append(dict)
                    for answer in answer_list:
                        try:
                            score = int(answer.answer)
                        except:
                            return JsonResponse({'status_code':5,'message':'得分不是整数'})
                        options[score-1]['choosed'] += 1
                    item['options'] = options

                item['options'] = options
                questions.append(item)
                print('success')
                response['questions'] = questions
            return JsonResponse(response)


        else:
            response = {'status_code': -1, 'message': 'invalid form'}
            return JsonResponse(response)
    else:
        response = {'status_code': -2, 'message': '请求错误'}
        return JsonResponse(response)
